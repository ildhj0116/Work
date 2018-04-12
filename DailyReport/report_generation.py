# -*- coding: utf-8 -*-
"""
Created on Tue Apr 10 08:25:48 2018

@author: Administrator
"""


from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
import pandas as pd

#from PIL import Image

def edit_picture(run,date,name,w):
    run.clear()
    path = u"output/Daily/" + date + "/" + name + ".jpg"
    run.add_picture(path,width=Inches(w))
    return run

def edit_text(p,index,text,date,Size,Bold):
    p.text = p.text[:(index+1)] + text
    p.style.font.size = Pt(Size)
    p.style.font.bold = Bold
    return p

def Report_Generation(report_date,head_list,tail_list,image_title_list):

    title_dict = {1:report_date}
    head_index_list = [12,16,20,25,31,42,45,49,52,56,59]
    tail_index_list = [13,17,21,26,32,43,46,50,53,57,60]
    image_index_list = [14,18,22,27,33,34,38,44,51,58,47,54,61,65,67,73,74,75,76]
    
    head_rank_dict = dict(zip(head_index_list,head_list))
    tail_rank_dict = dict(zip(tail_index_list,tail_list))
    image_name_dict = dict(zip(image_index_list,image_title_list))
    
    document = Document("report.docx")
    counter = 0
    for p in document.paragraphs:
        counter += 1
        if counter == 1:
            text = title_dict[counter] + u"量化日报"
            p = edit_text(p,-1,text,report_date,20,True)
        elif counter in head_index_list:
            index = p.text.find(':')
            text = u'、'.join(head_rank_dict[counter])
            p = edit_text(p,index,text,report_date,12,False)
        elif counter in tail_index_list:
            index = p.text.find(':')
            text = u'、'.join(tail_rank_dict[counter])
            p = edit_text(p,index,text,report_date,12,False)
        elif counter in image_index_list:
            r = p.runs
            if counter == 27 or counter == 33 or counter == 34 or counter == 65:
                r[0] = edit_picture(r[0],report_date,image_name_dict[counter],6.2)
            else:
                r[0] = edit_picture(r[0],report_date,image_name_dict[counter],6.8)

            
    
    
            
    document.styles['Normal'].font.name = u'楷体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')        
    document.save("report/Daily/" + report_date + u"_量化日报.docx")        
 



if __name__ == "__main__":
    report_date = "2018-04-09"
    title_list = ([u"品种日收益",u"品种周收益",u"品种月收益",u"品种日振幅",u"品种日增仓",u"品种日减仓",u"会员持仓占比",u"多头日占比变动",
                       u"多头周占比变动",u"多头月占比变动",u"空头日占比变动",u"空头周占比变动",u"空头月占比变动",u"板块沉淀资金",u"品种沉淀资金",
                       u"资金流向变动率",u"总资金流向",u"主动资金流向",u"被动资金流向"])
    head_df = pd.read_csv("output/Daily/" + report_date + '/' + "head.csv",encoding="utf_8_sig",index_col=0)
    tail_df = pd.read_csv("output/Daily/" + report_date + '/' + "tail.csv",encoding="utf_8_sig",index_col=0)
    column_name_list = [u"1日收益",u"5日收益",u"20日收益",u"振幅",u"成交量变化",u"日多头持仓占比变化",u"日空头持仓占比变化",
                      u"周多头持仓占比变化",u"周空头持仓占比变化",u"月多头持仓占比变化",u"月空头持仓占比变化"]
    head_value_list = head_df[column_name_list].T.values.tolist()
    tail_value_list = tail_df[column_name_list].T.values.tolist()
    Report_Generation(report_date,head_value_list,tail_value_list,title_list)





       
